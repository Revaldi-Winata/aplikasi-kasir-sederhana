import docx

def safe_set_text(cell, text):
    cell.text = ""
    lines = text.replace('\r', '').split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            cell.paragraphs[0].text = line
        else:
            cell.add_paragraph(line)

def replace_text_in_paragraphs(paragraphs, old_text, new_text):
    for p in paragraphs:
        if old_text in p.text:
            p.text = p.text.replace(old_text, new_text)

doc = docx.Document(r"E:\File Kuliah\Semester 6\Pemrograman II\Praktek\TokoBerkahJaya\doc_workspace\TokoBerkahJaya_Dokumentasi.docx")

replace_text_in_paragraphs(doc.paragraphs, "Aplikasi Untuk Sistem xxx", "Aplikasi Kasir Toko Berkah Jaya")
replace_text_in_paragraphs(doc.paragraphs, "sistem XXX", "Aplikasi Kasir Toko Berkah Jaya")
replace_text_in_paragraphs(doc.paragraphs, "Sistem XXX", "Aplikasi Kasir Toko Berkah Jaya")
replace_text_in_paragraphs(doc.paragraphs, "Implementasi Tabel xxx", "Implementasi Tabel Kategori, Barang, Customer, Penjualan, dll.")
replace_text_in_paragraphs(doc.paragraphs, "xxx", "Toko Berkah Jaya")

doc.tables[0].cell(3, 2).text = "Apache NetBeans IDE 25, Java 11"
doc.tables[0].cell(4, 2).text = "MySQL (XAMPP Server)"

sql_tables = '''CREATE TABLE 	buser (
  id_user int(11) NOT NULL AUTO_INCREMENT,
  username varchar(50) NOT NULL,
  password varchar(100) NOT NULL,
  level enum('Admin','Kasir') NOT NULL,
  PRIMARY KEY (id_user)
);

CREATE TABLE 	bbarang (
  id_barang int(11) NOT NULL AUTO_INCREMENT,
  
ama_barang varchar(100) NOT NULL,
  id_kategori int(11) NOT NULL,
  harga decimal(10,2) NOT NULL,
  stok int(11) NOT NULL,
  PRIMARY KEY (id_barang)
);

CREATE TABLE 	bpenjualan (
  
o_faktur varchar(20) NOT NULL,
  	anggal date NOT NULL,
  	otal_bayar decimal(10,2) NOT NULL,
  PRIMARY KEY (
o_faktur)
);'''
safe_set_text(doc.tables[3].cell(0, 0), sql_tables)

code_snippet = '''Nama Modul : Koneksi.java
Deskripsi  : Mengatur koneksi ke database MySQL menggunakan JDBC

package database;
import java.sql.Connection;
import java.sql.DriverManager;

public class Koneksi {
    private static Connection koneksi;
    public static Connection getKoneksi() {
        if (koneksi == null) {
            String url = "jdbc:mysql://localhost:3306/db_toko_berkah_jaya";
            koneksi = DriverManager.getConnection(url, "root", "");
        }
        return koneksi;
    }
}'''
safe_set_text(doc.tables[5].cell(0, 0), code_snippet)
safe_set_text(doc.tables[5].cell(1, 0), "Hasil:\nKoneksi berhasil dan aplikasi dapat mengambil data dari database.")

# Test Plan Table
for i in range(1, len(doc.tables[6].rows)):
    for j in range(3):
        doc.tables[6].cell(i, j).text = ""

doc.tables[6].cell(1, 0).text = "Login"
doc.tables[6].cell(1, 1).text = "Verifikasi akses dan level"
doc.tables[6].cell(1, 2).text = "Black box"

doc.tables[6].cell(2, 0).text = "Kelola Barang"
doc.tables[6].cell(2, 1).text = "Tambah, Edit, Hapus Barang"
doc.tables[6].cell(2, 2).text = "Black box"

doc.tables[6].cell(3, 0).text = "Kasir / Penjualan"
doc.tables[6].cell(3, 1).text = "Proses transaksi & hitung kembalian"
doc.tables[6].cell(3, 2).text = "Black box"

doc.tables[6].cell(4, 0).text = "Laporan"
doc.tables[6].cell(4, 1).text = "Export laporan ke Excel & filter tanggal"
doc.tables[6].cell(4, 2).text = "Black box"

doc.tables[7].cell(0, 0).text = "Deskripsi: Pengujian Transaksi Penjualan\nProsedur: Menambahkan barang ke keranjang dan memproses pembayaran\nData Masukkan: Barang yang dipilih, jumlah beli, dan uang tunai pelanggan"
doc.tables[7].cell(2, 1).text = "Barang masuk keranjang, total harga otomatis. Stok berkurang dan tersimpan ke DB."
doc.tables[7].cell(7, 1).text = "Pesan Peringatan: 'Uang tidak cukup!' atau 'Pilih barang terlebih dahulu'"

doc.save(r"E:\File Kuliah\Semester 6\Pemrograman II\Praktek\TokoBerkahJaya\TokoBerkahJaya-Dokumentasi.docx")
print("Done")
